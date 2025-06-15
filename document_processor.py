"""
Document processing module for the RAG application.
Handles loading, parsing, and chunking of various document formats.
"""

import os
import logging
from typing import List, Dict, Optional, Union
from pathlib import Path
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from utils import clean_text, get_file_hash, validate_file_type, validate_file_size
from config import Config

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document loading, parsing, and chunking."""
    
    def __init__(self):
        self.config = Config()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.CHUNK_SIZE,
            chunk_overlap=self.config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_document(self, file_path: Union[str, Path]) -> Dict[str, Union[str, Dict]]:
        """
        Load and parse a document from file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document content and metadata
        """
        file_path = Path(file_path)
        
        # Validate file
        if not self._validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        # Extract content based on file type
        content = ""
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == ".pdf":
                content = self._extract_pdf_content(file_path)
            elif file_extension == ".txt":
                content = self._extract_txt_content(file_path)
            elif file_extension == ".docx":
                content = self._extract_docx_content(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Clean the extracted content
            content = clean_text(content)
            
            # Generate metadata
            metadata = self._generate_metadata(file_path, content)
            
            return {
                "content": content,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    def chunk_document(self, document: Dict[str, Union[str, Dict]]) -> List[Dict[str, Union[str, Dict]]]:
        """
        Chunk a document into smaller pieces.
        
        Args:
            document: Document dictionary with content and metadata
            
        Returns:
            List of document chunks with metadata
        """
        content = document["content"]
        metadata = document["metadata"]
        
        if not content:
            logger.warning("Empty document content, skipping chunking")
            return []
        
        try:
            # Split the text into chunks
            chunks = self.text_splitter.split_text(content)
            
            # Create chunk documents with metadata
            chunk_documents = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_id": i,
                    "chunk_count": len(chunks),
                    "chunk_size": len(chunk)
                })
                
                chunk_documents.append({
                    "content": chunk,
                    "metadata": chunk_metadata
                })
            
            logger.info(f"Document chunked into {len(chunk_documents)} pieces")
            return chunk_documents
            
        except Exception as e:
            logger.error(f"Error chunking document: {e}")
            raise
    
    def process_document(self, file_path: Union[str, Path]) -> List[Dict[str, Union[str, Dict]]]:
        """
        Complete document processing pipeline: load and chunk.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of document chunks ready for embedding
        """
        try:
            # Load the document
            document = self.load_document(file_path)
            
            # Chunk the document
            chunks = self.chunk_document(document)
            
            logger.info(f"Successfully processed document: {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error in document processing pipeline for {file_path}: {e}")
            raise
    
    def _validate_file(self, file_path: Path) -> bool:
        """Validate file existence, type, and size."""
        if not file_path.exists():
            logger.error(f"File does not exist: {file_path}")
            return False
        
        if not validate_file_type(file_path, self.config.SUPPORTED_FILE_TYPES):
            logger.error(f"Unsupported file type: {file_path.suffix}")
            return False
        
        if not validate_file_size(file_path, self.config.MAX_FILE_SIZE_MB):
            logger.error(f"File too large: {file_path}")
            return False
        
        return True
    
    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text content from PDF file."""
        content = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1} from {file_path}: {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            raise
        
        return content
    
    def _extract_txt_content(self, file_path: Path) -> str:
        """Extract text content from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                return file.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            raise
    
    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract text content from DOCX file."""
        try:
            doc = docx.Document(file_path)
            content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            return content
            
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise
    
    def _generate_metadata(self, file_path: Path, content: str) -> Dict:
        """Generate metadata for a document."""
        try:
            file_stats = file_path.stat()
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": file_path.suffix.lower(),
                "file_size": file_stats.st_size,
                "created_time": file_stats.st_ctime,
                "modified_time": file_stats.st_mtime,
                "content_length": len(content),
                "file_hash": get_file_hash(file_path),
                "word_count": len(content.split()) if content else 0
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error generating metadata for {file_path}: {e}")
            return {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": file_path.suffix.lower(),
                "error": str(e)
            }

